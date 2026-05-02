import logging
from collections.abc import Callable
from pathlib import Path

from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import Document

from app.indexing.parser import parse_openapi_spec

logger = logging.getLogger(__name__)

LoaderFn = Callable[[Path], list[Document]]

_registry: dict[str, LoaderFn] = {}


def register_loader(*extensions: str) -> Callable[[LoaderFn], LoaderFn]:
    def decorator(fn: LoaderFn) -> LoaderFn:
        for ext in extensions:
            _registry[ext.lower().lstrip(".")] = fn
        return fn

    return decorator


def _make_metadata(path: Path, doc_type: str = "guide") -> dict[str, str]:
    return {"doc_type": doc_type, "source": str(path), "filename": path.name}


@register_loader("md", "txt", "rst")
def _load_text(path: Path) -> list[Document]:
    return [
        Document(
            text=path.read_text(encoding="utf-8"),
            metadata=_make_metadata(path),
        )
    ]


@register_loader("json")
def _load_openapi_json(path: Path) -> list[Document]:
    return parse_openapi_spec(path)


@register_loader("yaml", "yml")
def _load_yaml(path: Path) -> list[Document]:
    import yaml

    with open(path) as f:
        spec = yaml.safe_load(f)

    if isinstance(spec, dict) and "openapi" in spec:
        import json
        import tempfile

        with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as tmp:
            json.dump(spec, tmp, ensure_ascii=False)
            tmp_path = Path(tmp.name)
        try:
            return parse_openapi_spec(tmp_path)
        finally:
            tmp_path.unlink()

    return [
        Document(
            text=path.read_text(encoding="utf-8"),
            metadata=_make_metadata(path),
        )
    ]


@register_loader("pdf")
def _load_pdf(path: Path) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            pages.append(text)

    if not pages:
        return []

    return [
        Document(
            text="\n\n".join(pages),
            metadata=_make_metadata(path),
        )
    ]


@register_loader("docx")
def _load_docx(path: Path) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    if not paragraphs:
        return []

    return [
        Document(
            text="\n\n".join(paragraphs),
            metadata=_make_metadata(path),
        )
    ]


@register_loader("html", "htm")
def _load_html(path: Path) -> list[Document]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(path.read_text(encoding="utf-8"), "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    if not text:
        return []

    return [
        Document(
            text=text,
            metadata=_make_metadata(path),
        )
    ]


@register_loader("csv")
def _load_csv(path: Path) -> list[Document]:
    import csv

    with open(path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        rows = [" | ".join(f"{k}: {v}" for k, v in row.items() if v) for row in reader]

    if not rows:
        return []

    return [
        Document(
            text="\n".join(rows),
            metadata=_make_metadata(path),
        )
    ]


def load_documents(docs_dir: Path) -> list[Document]:
    documents: list[Document] = []

    for file_path in sorted(docs_dir.iterdir()):
        if file_path.is_dir() or file_path.name.startswith("."):
            continue

        ext = file_path.suffix.lower().lstrip(".")
        loader = _registry.get(ext)

        if loader is None:
            logger.warning("No loader for .%s — skipping %s", ext, file_path.name)
            continue

        try:
            docs = loader(file_path)
            documents.extend(docs)
            logger.info("Loaded %s (%d documents)", file_path.name, len(docs))
        except Exception:
            logger.exception("Failed to load %s", file_path.name)

    return documents


def get_supported_extensions() -> list[str]:
    return sorted(_registry.keys())


def get_splitter() -> SentenceSplitter:
    return SentenceSplitter(chunk_size=1024, chunk_overlap=200)
